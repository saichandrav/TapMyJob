import pdfplumber
import io
import re

def parse_resume_bytes(file_bytes: bytes, filename: str = "") -> str:
    """
    Extracts text from a resume file. Supports PDF and DOCX formats.
    """
    fname = (filename or "").lower()

    if fname.endswith(".docx") or fname.endswith(".doc"):
        return parse_docx_resume(file_bytes)
    else:
        # Default to PDF
        return parse_pdf_resume(file_bytes)


def parse_pdf_resume(file_bytes: bytes) -> str:
    """
    Extracts text from a PDF resume using pdfplumber.
    """
    try:
        text_content = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)

        full_text = "\n".join(text_content)

        if not full_text.strip():
            raise Exception("No text could be extracted from the PDF. It might be an image-based PDF.")

        return full_text.strip()

    except Exception as e:
        raise Exception(f"Error parsing PDF: {str(e)}")


def parse_docx_resume(file_bytes: bytes) -> str:
    """
    Extracts text from a DOCX resume using python-docx.
    """
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

        full_text = "\n".join(paragraphs)

        if not full_text.strip():
            raise Exception("No text could be extracted from the DOCX file.")

        return full_text.strip()

    except ImportError:
        raise Exception("python-docx is not installed. Run: pip install python-docx")
    except Exception as e:
        raise Exception(f"Error parsing DOCX: {str(e)}")

def extract_personal_info(resume_text: str) -> dict:
    """
    Extract personal information from resume text using regex patterns.
    """
    info = {}

    # Extract email
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    email_match = re.search(email_pattern, resume_text)
    if email_match:
        info['email'] = email_match.group()

    # Extract phone (Indian format)
    phone_pattern = r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    phone_match = re.search(phone_pattern, resume_text)
    if phone_match:
        info['phone'] = phone_match.group()

    # Extract LinkedIn URL
    linkedin_pattern = r'(https?://)?(www\.)?linkedin\.com/in/[a-zA-Z0-9-]+'
    linkedin_match = re.search(linkedin_pattern, resume_text)
    if linkedin_match:
        url = linkedin_match.group()
        if not url.startswith('http'):
            url = 'https://' + url
        info['linkedin'] = url

    # Extract GitHub URL
    github_pattern = r'(https?://)?(www\.)?github\.com/[a-zA-Z0-9-]+'
    github_match = re.search(github_pattern, resume_text)
    if github_match:
        url = github_match.group()
        if not url.startswith('http'):
            url = 'https://' + url
        info['github'] = url

    # Extract CodeChef URL
    codechef_pattern = r'(https?://)?(www\.)?codechef\.com/users/[a-zA-Z0-9-]+'
    codechef_match = re.search(codechef_pattern, resume_text)
    if codechef_match:
        url = codechef_match.group()
        if not url.startswith('http'):
            url = 'https://' + url
        info['codechef'] = url

    # Extract LeetCode URL
    leetcode_pattern = r'(https?://)?(www\.)?leetcode\.com/[a-zA-Z0-9/-]+'
    leetcode_match = re.search(leetcode_pattern, resume_text)
    if leetcode_match:
        url = leetcode_match.group()
        if not url.startswith('http'):
            url = 'https://' + url
        info['leetcode'] = url

    # Extract name (usually first line or in caps)
    lines = resume_text.split('\n')
    for line in lines[:5]:
        line = line.strip()
        if line and len(line) > 2 and len(line) < 50 and not re.search(r'[0-9@]', line):
            # Check if it looks like a name (no numbers, no special chars except spaces)
            if re.match(r'^[A-Z\s]+$', line):
                info['name'] = line
                break

    return info
